import streamlit as st
import google.generativeai as genai
from PIL import Image
import io
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
import docx

# --- CONFIGURAZIONE ---
st.set_page_config(page_title="InclusivAI-tool", layout="wide")

API_KEY = "AIzaSyBqab5Ib035fmA_n4mF1X_1ptCMWqmHabY"

genai.configure(api_key=API_KEY)
model = genai.GenerativeModel(model_name="gemini-2.5-flash")

# --- FUNZIONI DI ESPORTAZIONE ---

def crea_pdf(testo):
    """Genera un file PDF dal testo fornito."""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    c.setFont("Helvetica", 12)
    
    # Gestione del testo a capo automatico
    lines = simpleSplit(testo, "Helvetica", 12, width - 100)
    y = height - 50
    for line in lines:
        if y < 50: # Nuova pagina se finisce lo spazio
            c.showPage()
            c.setFont("Helvetica", 12)
            y = height - 50
        c.drawString(50, y, line)
        y -= 15
    c.save()
    buf.seek(0)
    return buf

def crea_word(testo):
    """Genera un file Word (.docx) dal testo fornito."""
    buf = io.BytesIO()
    doc = docx.Document()
    doc.add_heading('InclusivAI - Test di Verifica', 0)
    doc.add_paragraph(testo)
    doc.save(buf)
    buf.seek(0)
    return buf

# --- INTERFACCIA APP ---
st.title("🚀 InclusivAI-tool")
st.write("Trasforma contenuti didattici in materiali inclusivi (PDF/Word).")

# Inizializziamo la memoria di Streamlit per salvare il risultato
if "risultato_ia" not in st.session_state:
    st.session_state["risultato_ia"] = None
if "tipo_servizio" not in st.session_state:
    st.session_state["tipo_servizio"] = None

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("📁 Input")
    file_caricato = st.file_uploader("Carica una foto o un PDF della pagina:", type=["jpg", "png", "pdf"])
    testo_manuale = st.text_area("Oppure incolla il testo qui:", height=200)
    
    st.info("Scegli la trasformazione:")
    c1, c2 = st.columns(2)
    c3, c4 = st.columns(2)
    btn_dsa = c1.button("✨ DSA (-> PDF)")
    btn_nai = c2.button("🌍 NAI (-> PDF)")
    btn_mappa = c3.button("🗺️ MAPPA (-> PDF)")
    btn_test = c4.button("📝 TEST (-> Word)")

# Controllo dei click dei bottoni per attivare l'IA
azione_richiesta = btn_dsa or btn_nai or btn_mappa or btn_test

with col2:
    st.subheader("📄 Output")
    
    if azione_richiesta:
        if not testo_manuale and not file_caricato:
            st.error("Inserisci del testo o un file!")
        else:
            with st.spinner("Elaborazione in corso con Gemini..."):
                contenuto = []
                if file_caricato:
                    if file_caricato.type == "application/pdf":
                        contenuto.append({'mime_type': 'application/pdf', 'data': file_caricato.read()})
                    else:
                        contenuto.append(Image.open(file_caricato))
                
                testo_base = testo_manuale if testo_manuale else "Analizza il file allegato."
                
                # Definizione Prompt Rigidi in Italiano
                if btn_dsa:
                    prompt = f"Agisci come esperto DSA. Semplifica questo contenuto in italiano: {testo_base}. Usa frasi brevi, elenchi puntati e grassetti terapeutici."
                    st.session_state["tipo_servizio"] = "DSA"
                elif btn_nai:
                    prompt = f"Agisci come facilitatore L2 livello A2. Semplifica per alunni stranieri questo contenuto in italiano: {testo_base}."
                    st.session_state["tipo_servizio"] = "NAI"
                elif btn_mappa:
                    prompt = f"Crea una mappa concettuale testuale gerarchica e ben strutturata in italiano da questo contenuto: {testo_base}."
                    st.session_state["tipo_servizio"] = "MAPPA"
                elif btn_test:
                    prompt = f"Crea un test in italiano (5 domande a risposta multipla, 5 vero/falso) con soluzioni basato su questo contenuto: {testo_base}."
                    st.session_state["tipo_servizio"] = "TEST"
                
                contenuto.append(prompt)
                
                try:
                    response = model.generate_content(contenuto)
                    st.session_state["risultato_ia"] = response.text
                except Exception as e:
                    st.error(f"Errore durante la chiamata a Gemini: {e}")

    # Mostra il risultato se presente in memoria e abilita il download stabile
    if st.session_state["risultato_ia"]:
        st.markdown(st.session_state["risultato_ia"])
        
        if st.session_state["tipo_servizio"] == "TEST":
            file_word = crea_word(st.session_state["risultato_ia"])
            st.download_button("📥 Scarica TEST (.docx)", file_word, "test_inclusivai.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            file_pdf = crea_pdf(st.session_state["risultato_ia"])
            st.download_button("📥 Scarica DOCUMENTO (.pdf)", file_pdf, "inclusivai_output.pdf", mime="application/pdf")
        
